import pandas as pd
from docx import Document
from datetime import date, datetime
from dateutil.relativedelta import relativedelta
import ttkbootstrap as ttk
from ttkbootstrap.constants import * 
from tkinter import *
from ttkbootstrap.dialogs import Messagebox
import os

tabela_excel = "DADOS/controle.xlsx"
carteira_idoso = "DADOS/modelo_idoso.docx"
carteira_pne = "DADOS/modelo_pne.docx"
salvar_aqui = f"{datetime.now().year}"

def ultimo_cartao(df):
    cartao = int(df['Cartao n'].iloc[-1])
    anterior = int(df['Cartao n'].iloc[-2])
    data = df['Emissao'].iloc[-1]
    data = int(data[6:])
    hoje = datetime.now().year
    if(anterior>1 and hoje - data == 1):
        cartao = 0
    return cartao

def encontra_cartao(df):
    controle = df.loc[df['Solicitante']==nome.get()]
    if(len(controle)>0):
        controle = controle['Emissao'].iloc[-1]
        controle= str(controle)
        mb = Messagebox.show_question(f"Foi encontrado uma carteirinha em nome de {nome.get()} no dia {controle}, quer continuar?", "Aviso")
        return mb
        

def excel(nome, data_entrada, tipo, dia_emissao, dia_validade):
    df = pd.read_excel(tabela_excel)
    cartao = ultimo_cartao(df)
    nova_entrada = [nome, data_entrada, cartao+1, dia_emissao, tipo, dia_validade]
    df.loc[len(df)] = nova_entrada
    df.to_excel(tabela_excel, index=False)

def word(nome, tipo, dia_emissao, dia_validade):
    if(tipo=='IDOSO'):
        documento = Document(carteira_idoso)
    else:
        documento = Document(carteira_pne)
    df = pd.read_excel(tabela_excel)
    cartao = ultimo_cartao(df)
    if(ultimo_cartao(df)==0):
        cartao = 1
    else:
        cartao = ultimo_cartao(df)
    cartao = str(cartao)
    for paragrafo in documento.paragraphs:
        for run in paragrafo.runs:
            if("XXX" in run.text):
                run.text = run.text.replace("XXX", cartao)
            if("dd/mm/aa" in run.text):
                run.text = run.text.replace("dd/mm/aa", dia_validade)
            if("DD/MM/AA" in run.text):
                run.text = run.text.replace("DD/MM/AA", dia_emissao)
            if("SOLICITANTE" in run.text):
                run.text = run.text.replace("SOLICITANTE", nome)
    pasta_nova = f"{datetime.now().year}"
    if(not os.path.exists(pasta_nova)):
        os.mkdir(pasta_nova)
    documento.save(f"{salvar_aqui}/{cartao} - {nome}.docx")

def criar():
    if (len(nome.get())>0 and len(my_topping.get())>0):
        df = pd.read_excel(tabela_excel)
        if(encontra_cartao(df)!="Não"):
            tipo = my_topping.get()
            dia_emissao = date.today()
            dia_validade = dia_emissao + relativedelta(years=5)
            dia_emissao = dia_emissao.strftime('%d/%m/%Y')
            dia_validade = dia_validade.strftime('%d/%m/%Y')
            excel(nome.get(), data.entry.get(), tipo, dia_emissao, dia_validade)
            word(nome.get(), tipo, dia_emissao, dia_validade)
            mb = Messagebox.show_info("Carteirinha feita com sucesso!!", "Sucesso")
        else:
            mb = Messagebox.show_info("Carteirinha cancelada!!", "Aviso")
    else:
        if(len(nome.get())==0):
            mb = Messagebox.show_warning("Digite o nome do solicitante!!", "Aviso")
        if(len(my_topping.get())==0):
            mb = Messagebox.show_warning("Escolha o tipo da carteirinha!!", "Aviso")

root = ttk.Window(themename="darkly")
root.title("CRIADOR DE CARTEIRINHAS")
root.iconbitmap('DADOS/imagem.ico')
root.geometry("400x400") 

my_style = ttk.Style()
my_style.configure('success.Outline.TButton', font=("Helvetica", 10))

rotulo = ttk.Label(text="CRIADOR DE CARTEIRINHAS", font=("arial", 15), bootstyle=DEFAULT)
rotulo.pack(pady=25)

texto = ttk.Label(root, text="Digite o nome do solicitante:", bootstyle=INFO, font=("arial", 10))
texto.pack(pady=1, fill='x')
nome = ttk.Entry(root)
nome.pack(pady=10, fill='x')

texto = ttk.Label(root, text="Digite a data de entrada:", bootstyle=INFO, font=("arial", 10))
texto.pack(pady=1, fill='x')
data = ttk.DateEntry(root, bootstyle="primary")
data.pack(pady=10, anchor='w')

texto = ttk.Label(root, text="Escolha o tipo da carteirinha:", bootstyle=INFO, font=("arial", 10))
texto.pack(pady=1, fill='x')
toppings = ["IDOSO", "PNE"]
my_topping = StringVar()
for topping in toppings:
    ttk.Radiobutton(root, bootstyle="primary", variable=my_topping, text = topping, value = topping).pack(pady=5, fill='x')

botao_gerar = ttk.Button(text="GERAR", bootstyle=(SUCCESS), style=("success.Outline.TButton"), width=30, command=criar)
botao_gerar.pack(pady=25)

my_sizegrip = ttk.Sizegrip(root, bootstyle="primary")
my_sizegrip.pack(anchor="se", fill="both", expand=True)
root.mainloop()